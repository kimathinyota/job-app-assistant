import os
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional, Dict
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------
# HELPER: LaTeX Environment Setup
# ---------------------------------------------------------
def escape_latex(text: str) -> str:
    if not text: return ""
    chars = {
        "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#", "_": r"\_",
        "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}", "\\": r"\textbackslash{}",
    }
    return "".join(chars.get(c, c) for c in str(text))

def create_latex_env(template_dir: str):
    env = Environment(
        loader=FileSystemLoader(template_dir),
        block_start_string='\\BLOCK{', block_end_string='}',
        variable_start_string='\\VAR{', variable_end_string='}',
        comment_start_string='\\#{', comment_end_string='}',
        trim_blocks=True, autoescape=False,
    )
    env.filters['latex_escape'] = escape_latex 
    return env

# ---------------------------------------------------------
# PDF Generator (LaTeX)
# ---------------------------------------------------------
class PDFGenerator:
    LATEX_PATH = "/Library/TeX/texbin/pdflatex" # Change if on Linux/Windows

    def __init__(self, template_dir: str):
        self.env = create_latex_env(template_dir)

    def render_cv(self, context: dict, section_order: List[str] = None, section_titles: Dict[str, str] = None) -> bytes:
        if not section_order:
            section_order = ["summary", "education", "skills", "projects", "experience", "hobbies"]
        
        context['section_order'] = [s.lower() for s in section_order]
        context['section_titles'] = section_titles or {}

        template = self.env.get_template('cv_template.tex')
        tex_content = template.render(**context)
        return self._compile_tex(tex_content)

    def _compile_tex(self, tex_content: str) -> bytes:
        with tempfile.TemporaryDirectory() as temp_dir:
            tex_path = Path(temp_dir) / "resume.tex"
            tex_path.write_text(tex_content, encoding='utf-8')

            for _ in range(2):
                subprocess.run(
                    [self.LATEX_PATH, "-interaction=nonstopmode", "resume.tex"],
                    cwd=temp_dir,
                    stdout=subprocess.DEVNULL,
                    check=True
                )

            pdf_path = Path(temp_dir) / "resume.pdf"
            return pdf_path.read_bytes()

# ---------------------------------------------------------
# Word Generator (Docx)
# ---------------------------------------------------------
class WordGenerator:
    
    def create_docx(self, cv_data: dict, skill_groups: dict, section_order: List[str] = None, section_titles: Dict[str, str] = None) -> Path:
        doc = Document()
        self._setup_page_layout(doc)
        self._add_header(doc, cv_data)

        if not section_order:
            section_order = ["summary", "education", "skills", "projects", "experience", "hobbies"]
        
        titles = section_titles or {}
        default_titles = {
            "summary": "Professional Summary",
            "education": "Education",
            "skills": "Technical Skills",
            "projects": "Academic & Research Projects",
            "experience": "Experience",
            "hobbies": "Interests & Hobbies"
        }

        renderers = {
            "summary": lambda: self._add_summary(doc, cv_data, titles.get('summary', default_titles['summary'])),
            "education": lambda: self._add_education(doc, cv_data, titles.get('education', default_titles['education'])),
            "skills": lambda: self._add_skills(doc, skill_groups, titles.get('skills', default_titles['skills'])),
            "projects": lambda: self._add_projects(doc, cv_data, titles.get('projects', default_titles['projects'])),
            "experience": lambda: self._add_experience(doc, cv_data, titles.get('experience', default_titles['experience'])),
            "hobbies": lambda: self._add_hobbies(doc, cv_data, titles.get('hobbies', default_titles['hobbies'])),
        }

        for section in section_order:
            section_key = section.lower()
            if section_key in renderers:
                renderers[section_key]()

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return Path(tmp.name)

    # --- Internal Helpers ---

    def _setup_page_layout(self, doc):
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # --- FONT SETTING: ARIAL ---
        style = doc.styles['Normal']
        style.font.name = 'Arial' 
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.0

    def _add_section_header(self, doc, title):
        p = doc.add_paragraph()
        run = p.add_run(str(title).upper()) 
        run.font.name = 'Arial' # Explicitly set font
        run.font.size = Pt(14)
        run.bold = True
        
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    def _add_header(self, doc, cv_data):
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        full_name = f"{cv_data.get('first_name', '')} {cv_data.get('last_name', '')}"
        if cv_data.get('is_title_in_name') and cv_data.get('title'):
            full_name = f"{cv_data.get('title')} {full_name}"
            
        name_run = header.add_run(full_name)
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(24)
        name_run.bold = True
        
        info = cv_data.get('contact_info', {})
        parts = []
        for k, v in info.items():
            if v: parts.append(v)
            
        contact_line = " | ".join(parts)
        contact_run = header.add_run(f"\n{contact_line}")
        contact_run.font.name = 'Arial'
        contact_run.font.size = Pt(11)

    def _format_cell_p(self, cell):
        """Helper to remove extra spacing from table cells for tighter look."""
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    def _add_bullet(self, doc, text):
        """Helper to add properly indented bullet points."""
        bull = doc.add_paragraph(text, style='List Bullet')
        bull.style.font.name = 'Arial'
        bull.paragraph_format.left_indent = Inches(0.5)
        bull.paragraph_format.first_line_indent = Inches(-0.25)
        bull.paragraph_format.space_before = Pt(0) 
        bull.paragraph_format.space_after = Pt(0)
        bull.paragraph_format.line_spacing = 1.0 
        return bull

    def _add_item_spacer(self, doc):
        """
        Adds a small visual separator between items.
        """
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.line_spacing = 1.0
        run = spacer.add_run()
        run.font.size = Pt(3) 

    def _prepare_date(self, date_str):
        if not date_str: return ""
        return date_str.replace("--", "to")

    # --- SECTION RENDERERS ---

    def _add_summary(self, doc, cv_data, title):
        if not cv_data.get('summary'): return
        self._add_section_header(doc, title)
        p = doc.add_paragraph(cv_data['summary'])
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0

    def _add_education(self, doc, cv_data, title):
        if not cv_data.get('education'): return
        self._add_section_header(doc, title)
        
        items = cv_data['education']
        for i, edu in enumerate(items):
            if i > 0: self._add_item_spacer(doc)

            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            # Arial is wider, so we ensure the left column has enough space
            table.columns[0].width = Inches(6.0)
            table.columns[1].width = Inches(1.5)
            
            date_str = self._prepare_date(edu.get('formatted_date'))
            
            r1 = table.rows[0]
            r1.cells[0].text = f"{edu.get('degree', '')} {edu.get('field', '')}"
            r1.cells[0].paragraphs[0].runs[0].bold = True
            r1.cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            
            r1.cells[1].text = date_str
            r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r1.cells[1].paragraphs[0].runs[0].font.name = 'Arial'
            
            r2 = table.rows[1]
            r2.cells[0].text = edu.get('institution', '')
            r2.cells[0].paragraphs[0].runs[0].italic = True
            r2.cells[0].paragraphs[0].runs[0].font.name = 'Arial'

            self._format_cell_p(r1.cells[0])
            self._format_cell_p(r1.cells[1])
            self._format_cell_p(r2.cells[0])
            self._format_cell_p(r2.cells[1])

            for ach in edu.get('achievements', []):
                self._add_bullet(doc, ach['text'])

    def _add_skills(self, doc, skill_groups, title):
        if not skill_groups: return
        self._add_section_header(doc, title)
        
        for cat, skills in skill_groups.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            t = p.add_run(f"{cat}: ")
            t.bold = True
            t.font.name = 'Arial'
            
            r = p.add_run(", ".join(skills))
            r.font.name = 'Arial'

    def _add_projects(self, doc, cv_data, title):
        if not cv_data.get('projects'): return
        self._add_section_header(doc, title)
        
        items = cv_data['projects']
        for i, proj in enumerate(items):
            if i > 0: self._add_item_spacer(doc)

            table = doc.add_table(rows=1, cols=2)
            table.autofit = False 
            table.columns[0].width = Inches(6.0)
            table.columns[1].width = Inches(1.5)
            
            r1 = table.rows[0]
            
            c1_p = r1.cells[0].paragraphs[0]
            t_run = c1_p.add_run(proj.get('title', ''))
            t_run.bold = True
            t_run.font.name = 'Arial'
            
            context = proj.get('context_display', '')
            if context:
                c_run = c1_p.add_run(f" | {context}")
                c_run.italic = True
                c_run.font.name = 'Arial'
            
            date_str = self._prepare_date(proj.get('formatted_date'))
            r1.cells[1].text = date_str
            r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r1.cells[1].paragraphs[0].runs[0].font.name = 'Arial'

            self._format_cell_p(r1.cells[0])
            self._format_cell_p(r1.cells[1])

            for ach in proj.get('achievements', []):
                self._add_bullet(doc, ach['text'])

    def _add_experience(self, doc, cv_data, title):
        if not cv_data.get('experiences'): return
        self._add_section_header(doc, title)
        
        items = cv_data['experiences']
        for i, exp in enumerate(items):
            if i > 0: self._add_item_spacer(doc)

            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(6.0)
            table.columns[1].width = Inches(1.5)
            
            date_str = self._prepare_date(exp.get('formatted_date'))
            
            r1 = table.rows[0]
            r1.cells[0].text = exp.get('title', '')
            r1.cells[0].paragraphs[0].runs[0].bold = True
            r1.cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            
            if exp.get('location'):
                r1.cells[1].text = exp.get('location', '')
                r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r1.cells[1].paragraphs[0].runs[0].font.name = 'Arial'
            
            r2 = table.rows[1]
            r2.cells[0].text = exp.get('company', '')
            r2.cells[0].paragraphs[0].runs[0].italic = True
            r2.cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            
            r2.cells[1].text = date_str
            r2.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2.cells[1].paragraphs[0].runs[0].font.name = 'Arial'
            
            self._format_cell_p(r1.cells[0])
            self._format_cell_p(r1.cells[1])
            self._format_cell_p(r2.cells[0])
            self._format_cell_p(r2.cells[1])

            for ach in exp.get('achievements', []):
                self._add_bullet(doc, ach['text'])

    def _add_hobbies(self, doc, cv_data, title):
        if not cv_data.get('hobbies'): return
        self._add_section_header(doc, title)
        
        items = cv_data['hobbies']
        for i, hobby in enumerate(items):
            if i > 0: self._add_item_spacer(doc)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(hobby.get('name', ''))
            run.bold = True
            run.font.name = 'Arial'

            for ach in hobby.get('achievements', []):
                self._add_bullet(doc, ach['text'])